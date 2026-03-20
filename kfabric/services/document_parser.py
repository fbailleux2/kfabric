from __future__ import annotations

import io
import re
from hashlib import sha256

from bs4 import BeautifulSoup

from kfabric.infra.models import CandidateDocument
from kfabric.services.content_payloads import BinaryPayload, unpack_binary_payload

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover - optional dependency
    DocxDocument = None

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover - optional dependency
    PdfReader = None

try:
    import trafilatura
except Exception:  # pragma: no cover - optional dependency
    trafilatura = None

try:
    from readability import Document as ReadabilityDocument
except Exception:  # pragma: no cover - optional dependency
    ReadabilityDocument = None


def parse_document(raw_content: str, content_type: str, candidate: CandidateDocument) -> dict[str, object]:
    binary_payload = unpack_binary_payload(raw_content)
    if binary_payload:
        extracted_title, text, headings, extraction_method = _parse_binary_payload(binary_payload, candidate)
    elif "markdown" in content_type or _looks_like_markdown(raw_content):
        extracted_title, text, headings, extraction_method = _parse_markdown_document(raw_content, candidate)
    elif "html" in content_type or _looks_like_html(raw_content):
        extracted_title, text, headings, extraction_method = _parse_html_document(raw_content, candidate)
    else:
        extracted_title, text, headings, extraction_method = _parse_text_document(raw_content, candidate)

    normalized_text = _normalize_text(text)
    if not normalized_text:
        normalized_text = candidate.snippet or candidate.title
    return {
        "normalized_text": normalized_text,
        "normalized_hash": sha256(normalized_text.encode("utf-8")).hexdigest(),
        "extracted_title": extracted_title,
        "headings": headings,
        "text_length": len(normalized_text),
        "extraction_method": extraction_method,
    }


def _extract_markdown_headings(raw_content: str) -> list[str]:
    headings: list[str] = []
    for line in raw_content.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            headings.append(stripped.lstrip("# ").strip())
    return headings[:12]


def _parse_binary_payload(
    payload: BinaryPayload,
    candidate: CandidateDocument,
) -> tuple[str, str, list[str], str]:
    if "pdf" in payload.media_type:
        return _parse_pdf_document(payload.data, candidate, payload.title)
    if "wordprocessingml" in payload.media_type or "docx" in payload.media_type:
        return _parse_docx_document(payload.data, candidate, payload.title)
    text = payload.data.decode("utf-8", errors="ignore")
    return candidate.title, text, _extract_markdown_headings(text), "binary_text"


def _parse_pdf_document(
    raw_pdf: bytes,
    candidate: CandidateDocument,
    payload_title: str | None,
) -> tuple[str, str, list[str], str]:
    extracted_title = payload_title or candidate.title
    if PdfReader is not None:
        try:
            reader = PdfReader(io.BytesIO(raw_pdf))
            metadata_title = _extract_pdf_metadata_title(reader)
            if metadata_title:
                extracted_title = metadata_title
            text = "\n\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
            if text:
                return extracted_title, text, _extract_candidate_headings(text), "pypdf"
        except Exception:
            pass

    literal_text = _extract_pdf_literal_text(raw_pdf)
    return extracted_title, literal_text or candidate.snippet or candidate.title, _extract_candidate_headings(literal_text), "pdf_literal_fallback"


def _extract_pdf_metadata_title(reader: PdfReader) -> str | None:
    metadata = reader.metadata
    if metadata is None:
        return None
    if hasattr(metadata, "title") and metadata.title:
        return str(metadata.title)
    if isinstance(metadata, dict):
        title = metadata.get("/Title") or metadata.get("title")
        if title:
            return str(title)
    return None


def _parse_docx_document(
    raw_docx: bytes,
    candidate: CandidateDocument,
    payload_title: str | None,
) -> tuple[str, str, list[str], str]:
    extracted_title = payload_title or candidate.title
    if DocxDocument is not None:
        try:
            document = DocxDocument(io.BytesIO(raw_docx))
            core_title = (document.core_properties.title or "").strip()
            if core_title:
                extracted_title = core_title
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            text = "\n".join(paragraphs).strip()
            if text:
                return extracted_title, text, _extract_candidate_headings(text), "docx"
        except Exception:
            pass

    fallback_text = candidate.snippet or candidate.title
    return extracted_title, fallback_text, _extract_candidate_headings(fallback_text), "docx_fallback"


def _extract_pdf_literal_text(raw_pdf: bytes) -> str:
    decoded = raw_pdf.decode("latin-1", errors="ignore")
    literal_matches = re.findall(r"\(((?:\\.|[^\\()]){3,})\)\s*T[Jj]", decoded)
    if not literal_matches:
        literal_matches = re.findall(r"\(((?:\\.|[^\\()]){3,})\)", decoded)

    cleaned_fragments: list[str] = []
    for match in literal_matches:
        fragment = match.strip()
        fragment = (
            fragment.replace(r"\(", "(")
            .replace(r"\)", ")")
            .replace(r"\n", "\n")
            .replace(r"\r", "")
            .replace(r"\t", "\t")
            .replace(r"\\", "\\")
        )
        fragment = re.sub(r"\s{2,}", " ", fragment).strip()
        if fragment and fragment not in cleaned_fragments:
            cleaned_fragments.append(fragment)
    return "\n".join(cleaned_fragments[:24])


def _parse_html_document(raw_content: str, candidate: CandidateDocument) -> tuple[str, str, list[str], str]:
    soup = BeautifulSoup(raw_content, "html.parser")
    extracted_title = soup.title.string.strip() if soup.title and soup.title.string else candidate.title
    headings = _extract_html_headings(soup)

    source_specific = _parse_source_specific_html(soup, candidate, extracted_title)
    if source_specific is not None:
        return source_specific

    if trafilatura is not None:
        try:
            extracted = trafilatura.extract(
                raw_content,
                include_comments=False,
                include_links=False,
                include_images=False,
                favor_precision=True,
                target_language=candidate.language or None,
            )
            if extracted:
                return extracted_title, extracted, headings or _extract_candidate_headings(extracted), "trafilatura"
        except Exception:
            pass

    if ReadabilityDocument is not None:
        try:
            readable = ReadabilityDocument(raw_content)
            summary = readable.summary()
            readable_soup = BeautifulSoup(summary, "html.parser")
            readable_text = readable_soup.get_text("\n", strip=True)
            if readable.short_title():
                extracted_title = readable.short_title()
            if readable_text:
                return extracted_title, readable_text, headings or _extract_candidate_headings(readable_text), "readability"
        except Exception:
            pass

    for tag in soup.select("script, style, noscript, svg, form, nav, footer"):
        tag.decompose()
    scope = soup.select_one("article, main, [role='main'], .article, .content, #content") or soup.body or soup
    text = scope.get_text("\n", strip=True)
    return extracted_title, text, headings or _extract_candidate_headings(text), "beautifulsoup"


def _parse_text_document(raw_content: str, candidate: CandidateDocument) -> tuple[str, str, list[str], str]:
    return candidate.title, raw_content, _extract_markdown_headings(raw_content) or _extract_candidate_headings(raw_content), "plain_text"


def _parse_markdown_document(raw_content: str, candidate: CandidateDocument) -> tuple[str, str, list[str], str]:
    headings = _extract_markdown_headings(raw_content)
    extracted_title = headings[0] if headings else candidate.title
    return extracted_title, raw_content, headings or _extract_candidate_headings(raw_content), "markdown"


def _extract_html_headings(soup: BeautifulSoup) -> list[str]:
    headings: list[str] = []
    for tag in soup.find_all(["h1", "h2", "h3"]):
        text = tag.get_text(" ", strip=True)
        if text and text not in headings:
            headings.append(text)
    return headings[:12]


def _extract_candidate_headings(raw_content: str) -> list[str]:
    headings: list[str] = []
    for line in raw_content.splitlines():
        cleaned = line.strip(" -*\t")
        if not cleaned:
            continue
        if cleaned.endswith(":"):
            headings.append(cleaned.rstrip(":"))
        elif cleaned.isupper() and len(cleaned) <= 80:
            headings.append(cleaned.title())
        elif len(cleaned.split()) <= 8 and 8 <= len(cleaned) <= 80 and cleaned[:1].isupper():
            headings.append(cleaned)
        if len(headings) >= 12:
            break
    deduplicated: list[str] = []
    for heading in headings:
        if heading not in deduplicated:
            deduplicated.append(heading)
    return deduplicated[:12]


def _normalize_text(text: str) -> str:
    normalized_lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    filtered_lines = [line for line in normalized_lines if line]
    normalized_text = "\n".join(filtered_lines)
    return re.sub(r"\n{3,}", "\n\n", normalized_text).strip()


def _looks_like_html(raw_content: str) -> bool:
    sample = raw_content[:512].lower()
    return "<html" in sample or "<body" in sample or "<article" in sample or "<main" in sample


def _looks_like_markdown(raw_content: str) -> bool:
    lines = [line.strip() for line in raw_content.splitlines()[:24] if line.strip()]
    if not lines:
        return False
    markdown_markers = sum(
        1
        for line in lines
        if line.startswith("#")
        or line.startswith(("- ", "* "))
        or line.startswith("```")
        or ("[" in line and "](" in line)
    )
    return markdown_markers >= 2


def _parse_source_specific_html(
    soup: BeautifulSoup,
    candidate: CandidateDocument,
    fallback_title: str,
) -> tuple[str, str, list[str], str] | None:
    if candidate.domain.lower() == "arxiv.org":
        return _parse_arxiv_html(soup, fallback_title)
    domain = candidate.domain.lower()
    if domain == "eur-lex.europa.eu":
        return _parse_eurlex_html(soup, fallback_title)
    if domain == "data.gouv.fr":
        return _parse_datagouv_html(soup, fallback_title)
    if domain == "legifrance.gouv.fr":
        return _parse_legifrance_html(soup, fallback_title)
    if domain == "hal.science":
        return _parse_hal_html(soup, fallback_title)
    if domain == "service-public.fr":
        return _parse_servicepublic_html(soup, fallback_title)
    if domain == "github.com":
        return _parse_github_html(soup, fallback_title)
    return None


def _parse_arxiv_html(soup: BeautifulSoup, fallback_title: str) -> tuple[str, str, list[str], str] | None:
    title = _text_from_selectors(
        soup,
        (
            "h1.title",
            ".title.mathjax",
            "main h1",
            "h1",
        ),
    ) or fallback_title
    scope = soup.select_one("main, article, #abs")
    if scope is None:
        return None

    authors = _text_from_selectors(scope, (".authors", ".authors a", "p"))
    abstract = _text_from_selectors(scope, (".abstract", "blockquote.abstract", "p"))
    text_parts = [title]
    if authors:
        text_parts.append("Auteurs : " + authors.replace("Authors:", "").strip())
    if abstract:
        text_parts.append("Resume : " + abstract.replace("Abstract:", "").strip())
    text_parts.append(scope.get_text("\n", strip=True))
    headings = _extract_html_headings(scope if isinstance(scope, BeautifulSoup) else soup)
    return title, "\n\n".join(part for part in text_parts if part), headings, "arxiv_html"


def _parse_github_html(soup: BeautifulSoup, fallback_title: str) -> tuple[str, str, list[str], str] | None:
    title = _text_from_selectors(
        soup,
        (
            "strong[itemprop='name'] a",
            "meta[property='og:title']",
            "main h1",
            "h1",
        ),
    ) or fallback_title
    scope = soup.select_one("article.markdown-body, #readme, main")
    if scope is None:
        return None

    lead = _text_from_selectors(scope, ("p", "article.markdown-body p"))
    useful_links = _collect_link_labels(scope, ("article.markdown-body a[href]", "#readme a[href]", "a[href]"))
    text_parts = [title]
    if lead:
        text_parts.append(lead)
    if useful_links:
        text_parts.append("Liens utiles : " + ", ".join(useful_links[:6]))
    text_parts.append(scope.get_text("\n", strip=True))
    headings = _extract_html_headings(scope if isinstance(scope, BeautifulSoup) else soup)
    return title, "\n\n".join(part for part in text_parts if part), headings, "github_html"


def _parse_eurlex_html(soup: BeautifulSoup, fallback_title: str) -> tuple[str, str, list[str], str] | None:
    title = _text_from_selectors(
        soup,
        (
            ".eli-main-title",
            ".oj-doc-ti",
            ".doc-ti",
            "h1",
        ),
    ) or fallback_title
    scope = soup.select_one("#text, #docHtml, main, article")
    if scope is None:
        return None

    metadata = _collect_metadata_pairs(
        soup,
        (".eli-data dt", ".eli-metadata dt", "dl dt"),
        (".eli-data dd", ".eli-metadata dd", "dl dd"),
    )
    text_parts = [title, *metadata]
    body_text = scope.get_text("\n", strip=True)
    if body_text:
        text_parts.append(body_text)
    headings = _extract_html_headings(scope if isinstance(scope, BeautifulSoup) else soup)
    return title, "\n\n".join(part for part in text_parts if part), headings, "eurlex_html"


def _parse_datagouv_html(soup: BeautifulSoup, fallback_title: str) -> tuple[str, str, list[str], str] | None:
    title = _text_from_selectors(
        soup,
        (
            "[data-testid='dataset-title']",
            ".fr-card__title",
            "main h1",
            "h1",
        ),
    ) or fallback_title
    scope = soup.select_one("main, article, .fr-container")
    if scope is None:
        return None

    lead = _text_from_selectors(scope, (".fr-text--lead", ".dataset-description", "p"))
    resource_labels = _collect_link_labels(scope, (".resource-card a[href]", ".fr-card a[href]", ".fr-btn[href]"))
    text_parts = [title]
    if lead:
        text_parts.append(lead)
    if resource_labels:
        text_parts.append("Ressources : " + ", ".join(resource_labels[:6]))
    text_parts.append(scope.get_text("\n", strip=True))
    headings = _extract_html_headings(scope if isinstance(scope, BeautifulSoup) else soup)
    return title, "\n\n".join(part for part in text_parts if part), headings, "datagouv_html"


def _parse_hal_html(soup: BeautifulSoup, fallback_title: str) -> tuple[str, str, list[str], str] | None:
    title = _text_from_selectors(
        soup,
        (
            ".result-title",
            ".Page-title",
            "main h1",
            "h1",
        ),
    ) or fallback_title
    scope = soup.select_one("main, article, .content, .Page-content")
    if scope is None:
        return None

    abstract = _text_from_selectors(scope, (".abstract", "#abstract", ".description", "p"))
    authors = _text_from_selectors(scope, (".authors", ".result-authors", ".notice-authors"))
    text_parts = [title]
    if authors:
        text_parts.append("Auteurs : " + authors)
    if abstract:
        text_parts.append("Resume : " + abstract)
    text_parts.append(scope.get_text("\n", strip=True))
    headings = _extract_html_headings(scope if isinstance(scope, BeautifulSoup) else soup)
    return title, "\n\n".join(part for part in text_parts if part), headings, "hal_html"


def _parse_legifrance_html(soup: BeautifulSoup, fallback_title: str) -> tuple[str, str, list[str], str] | None:
    title = _text_from_selectors(
        soup,
        (
            ".title-page",
            ".titreTexte",
            "main h1",
            "h1",
        ),
    ) or fallback_title
    scope = soup.select_one("main, article, .texte, .content")
    if scope is None:
        return None

    metadata = _collect_metadata_pairs(
        scope,
        (".list-inline dt", ".metadata dt", "dl dt"),
        (".list-inline dd", ".metadata dd", "dl dd"),
    )
    body = _text_from_selectors(scope, (".article-body", ".content-text", ".texte", "p"))
    text_parts = [title, *metadata]
    if body:
        text_parts.append(body)
    text_parts.append(scope.get_text("\n", strip=True))
    headings = _extract_html_headings(scope if isinstance(scope, BeautifulSoup) else soup)
    return title, "\n\n".join(part for part in text_parts if part), headings, "legifrance_html"


def _parse_servicepublic_html(soup: BeautifulSoup, fallback_title: str) -> tuple[str, str, list[str], str] | None:
    title = _text_from_selectors(
        soup,
        (
            ".page-title",
            ".fr-h1",
            "main h1",
            "h1",
        ),
    ) or fallback_title
    scope = soup.select_one("main, article, .sp-content, .fr-container")
    if scope is None:
        return None

    lead = _text_from_selectors(scope, (".introduction", ".fr-text--lead", ".chapo", "p"))
    useful_links = _collect_link_labels(scope, (".fr-callout a[href]", ".fr-card a[href]", ".fr-link[href]"))
    text_parts = [title]
    if lead:
        text_parts.append(lead)
    if useful_links:
        text_parts.append("Ressources utiles : " + ", ".join(useful_links[:6]))
    text_parts.append(scope.get_text("\n", strip=True))
    headings = _extract_html_headings(scope if isinstance(scope, BeautifulSoup) else soup)
    return title, "\n\n".join(part for part in text_parts if part), headings, "servicepublic_html"


def _text_from_selectors(node: BeautifulSoup, selectors: tuple[str, ...]) -> str:
    for selector in selectors:
        tag = node.select_one(selector)
        if tag:
            text = tag.get_text(" ", strip=True)
            if text:
                return text
    return ""


def _collect_metadata_pairs(
    node: BeautifulSoup,
    dt_selectors: tuple[str, ...],
    dd_selectors: tuple[str, ...],
) -> list[str]:
    dt_tags = _select_many(node, dt_selectors)
    dd_tags = _select_many(node, dd_selectors)
    metadata: list[str] = []
    for dt_tag, dd_tag in zip(dt_tags, dd_tags, strict=False):
        dt_text = dt_tag.get_text(" ", strip=True)
        dd_text = dd_tag.get_text(" ", strip=True)
        if dt_text and dd_text:
            metadata.append(f"{dt_text}: {dd_text}")
    return metadata[:8]


def _collect_link_labels(node: BeautifulSoup, selectors: tuple[str, ...]) -> list[str]:
    labels: list[str] = []
    for selector in selectors:
        for tag in node.select(selector):
            text = tag.get_text(" ", strip=True)
            if text and text not in labels:
                labels.append(text)
    return labels


def _select_many(node: BeautifulSoup, selectors: tuple[str, ...]) -> list[BeautifulSoup]:
    selected: list[BeautifulSoup] = []
    for selector in selectors:
        selected.extend(node.select(selector))
    return selected
